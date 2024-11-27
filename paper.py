
# conda activate torch107
# pip install python-docx pdfplumber

import docx
import pdfplumber
import re
import pandas as pd

def extract_column_from_docx(docx_path, column_name):
    """从 Word 文档中提取指定列内容"""
    doc = docx.Document(docx_path)
    column_data = []
    for table in doc.tables:
        # 假设表头在第一行
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if column_name in headers:
            column_index = headers.index(column_name)
            for row in table.rows[1:]:
                cell_text = row.cells[column_index].text.strip()
                if cell_text:
                    column_data.append(cell_text)
    return column_data

def extract_chinese_journals_from_pdf(pdf_path):
    """从 PDF 表格中提取中文期刊名称"""
    chinese_journals = []
    
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            # 提取页面中的表格
            table = page.extract_table()
            if table:
                # 遍历表格的每一行
                for row in table[1:]:  # 跳过表头
                    # 假设期刊名称在表格中的某一列（例如第2列）
                    journal_name = row[1]  # 根据实际表格结构调整索引
                    if journal_name and re.search(r'[\u4e00-\u9fa5]', journal_name):  # 检查是否包含中文字符
                        chinese_journals.append(journal_name.strip())
    
    return chinese_journals


def extract_text_from_pdf(pdf_path):
    """逐行读取 PDF 文件并存储为列表"""
    text_data = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:# 将每一页的文本按行拆分并添加到 text_data 列表
                lines = text.split('\n')  # 按行拆分文本
                lines = lines[1:-1]  # 去掉首尾
                text_data.extend(lines)  # 将每行添加到总列表中
    return text_data

def extract_journal_names(text_data):
    """从文本中提取期刊名称"""
    journal_names = []
    pattern = re.compile(r'\d+\s+[A-Za-z0-9]+\s+([^\d,]+)')  # 匹配期刊名称
    for page_text in text_data:
        matches = pattern.findall(page_text)
        for match in matches:
            journal_names.append(match.split())
    return journal_names

# 1. 扁平化嵌套列表，并将英文期刊合并为一个完整的英文名称
def flatten_and_merge(original_list):
    flattened_list = []
    for item in original_list:
        # 如果是英文期刊名称的部分列表，合并成一个字符串
        if isinstance(item, list) and all(isinstance(x, str) for x in item):
            flattened_list.append(' '.join(item))  # 合并为一个字符串
        else:
            flattened_list.append(item[0])  # 保留中文期刊名称
        
    return flattened_list


def compare_columns_2(docx_column, pdf_column):
    # 确保 pdf_column 中的每个元素都是字符串
    pdf_column = [str(item) for item in pdf_column]
    """对比两个列的数据"""
    set_docx = set(docx_column)
    set_pdf = set(pdf_column)
    return {
        "intersection": set_docx & set_pdf,  # 交集
        "only_in_docx": set_docx - set_pdf,  # Word 独有
        "only_in_pdf": set_pdf - set_docx   # PDF 独有
    }

def compare_columns_3(docx_column, pdf_column1, pdf_column2):
    """对比三个列的数据，返回在 pdf_column1 中但不在 pdf_column2 中的元素"""
    
    # 确保 pdf_column 中的每个元素都是字符串
    pdf_column1 = [str(item) for item in pdf_column1]
    pdf_column2 = [str(item) for item in pdf_column2]
    
    # 将 docx_column 和 pdf_column 转换为集合
    set_docx = set(docx_column)
    set_pdf1 = set(pdf_column1)
    set_pdf2 = set(pdf_column2)

    # 计算在 pdf_column1 中但不在 pdf_column2 中的 docx_column 元素
    docx_in_pdf1_not_pdf2 = set_docx & set_pdf1 - set_pdf2
    
    return docx_in_pdf1_not_pdf2

# 文件路径
docx_path = "/Users/rl/Documents/PhD_student/Untitled_Folder/knowledge/oil.docx"
pdf_path = "/Users/rl/Documents/PhD_student/Untitled_Folder/knowledge/keji.pdf"
# 示例 PDF 文件路径
beidapdf_path = "/Users/rl/Documents/PhD_student/Untitled_Folder/knowledge/beida.pdf"



column_name = "期刊名称"

# 提取 油田 word  数据
docx_list = extract_column_from_docx(docx_path, column_name)

# 提取 科技核心 中的文本
text_data = extract_text_from_pdf(pdf_path)
pdf_column = extract_journal_names(text_data)

# 提取北大 中文期刊名称
chinese_journals = extract_chinese_journals_from_pdf(beidapdf_path)

# print('油田核心',docx_list)
print('油田核心 共 ',len(docx_list))
pdf_list = flatten_and_merge(pdf_column)
# print('科技核心',pdf_list)
print('科技核心 共 ',len(pdf_list))
print('北大核心 共 ',len(chinese_journals))

# 对比结果
result = compare_columns_3(docx_list, pdf_list, chinese_journals)


# 输出结果
# print("交集 (共有):", comparison_result["intersection"])
# print("仅在 Word 文档中:", comparison_result["only_in_docx"])
# print("仅在 PDF 文档中:", comparison_result["only_in_pdf"])

# 将数据转换为 DataFrame
intersection_df = pd.DataFrame(list(result), columns=["油田核心&科技核心!北大核心"])

# 保存为 Excel 文件
intersection_df.to_excel("result.xlsx", index=False)
print("交集已保存到 'result.xlsx' 文件中。")



'''
必须得用word另存为docx才可以
'''
# import os

# docx_path = "/Users/rl/Documents/PhD_student/Untitled_Folder/knowledge/oil.docx"

# # 检查文件是否存在
# if not os.path.exists(docx_path):
#     print(f"File not found: {docx_path}")
# else:
#     print("File exists. Verifying content...")
#     try:
#         from docx import Document
#         document = Document(docx_path)
#         print("Word document loaded successfully!")
#     except Exception as e:
#         print(f"Error loading document: {e}")

# import zipfile

# if not zipfile.is_zipfile(docx_path):
#     print("Error: File is not a valid .docx file.")
# else:
#     print("File is a valid .docx file.")
